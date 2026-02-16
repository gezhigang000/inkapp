#!/usr/bin/env python3
"""
Agent 核心循环：多轮 LLM function-calling 工具调用。

适配 sidecar stdin/stdout 同步模式，复用现有 search_adapter 和 llm_adapter 的
provider→endpoint/key/model 映射。
"""

import json
import logging
import os
import re
import subprocess
import sys
import threading
import time
from datetime import datetime

logger = logging.getLogger("ink.agent")

# ---------------------------------------------------------------------------
# Provider → endpoint 映射（与 llm_adapter.py 保持一致）
# ---------------------------------------------------------------------------

PROVIDER_ENDPOINTS = {
    "deepseek": "https://api.deepseek.com/v1/chat/completions",
    "openai":   "https://api.openai.com/v1/chat/completions",
    "glm":      "https://open.bigmodel.cn/api/paas/v4/chat/completions",
    "doubao":   "https://ark.cn-beijing.volces.com/api/v3/chat/completions",
    "kimi":     "https://api.moonshot.cn/v1/chat/completions",
}

PROVIDER_KEY_NAMES = {
    "deepseek": "DEEPSEEK_API_KEY",
    "openai":   "OPENAI_API_KEY",
    "glm":      "GLM_API_KEY",
    "doubao":   "DOUBAO_API_KEY",
    "kimi":     "KIMI_API_KEY",
}

PROVIDER_MODEL_NAMES = {
    "deepseek": ("DEEPSEEK_MODEL", "deepseek-chat"),
    "openai":   ("OPENAI_MODEL",   "gpt-4o"),
    "glm":      ("GLM_MODEL",      "glm-4-flash"),
    "doubao":   ("DOUBAO_MODEL",    "doubao-1.5-pro-32k"),
    "kimi":     ("KIMI_MODEL",      "moonshot-v1-8k"),
}

# ---------------------------------------------------------------------------
# Tool definitions (OpenAI function-calling format)
# ---------------------------------------------------------------------------

TOOL_DEFINITIONS = [
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "Search the web for up-to-date information. "
                "Returns results with title, URL, and content snippet."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query string.",
                    }
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_python",
            "description": (
                "Execute Python code in the workspace. "
                "Pre-installed libraries: python-docx, openpyxl, reportlab, "
                "pandas, numpy, matplotlib, Pillow. "
                "Timeout 30s. Returns stdout+stderr."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {
                        "type": "string",
                        "description": "The Python code to execute.",
                    }
                },
                "required": ["code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": (
                "Read a file from the workspace. "
                "Path is relative to workspace root, e.g. 'input/data.xlsx'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Relative path within the workspace.",
                    }
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": (
                "Write content to a file in the workspace. "
                "Creates parent directories if needed. "
                "You MUST use this to create 'output/article.html'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Relative path within the workspace.",
                    },
                    "content": {
                        "type": "string",
                        "description": "The content to write.",
                    },
                },
                "required": ["path", "content"],
            },
        },
    },
]


# ---------------------------------------------------------------------------
# Workspace management
# ---------------------------------------------------------------------------

def init_workspace(task_id, base_dir=None):
    """Create isolated workspace: {base_dir}/{task_id}/{input,data,output}/"""
    if base_dir is None:
        if sys.platform == "win32":
            _app_data = os.environ.get("APPDATA", os.path.expanduser("~"))
            base_dir = os.path.join(_app_data, "Ink", "agent-workspace")
        else:
            base_dir = os.path.join(os.path.expanduser("~"), ".ink", "agent-workspace")

    workspace = os.path.join(base_dir, task_id)
    for sub in ("input", "data", "output"):
        os.makedirs(os.path.join(workspace, sub), exist_ok=True)
    return workspace


def validate_path(path, workspace):
    """Prevent directory traversal. Returns absolute path."""
    resolved = os.path.normpath(os.path.join(workspace, path))
    if not resolved.startswith(os.path.normpath(workspace)):
        raise ValueError(f"Path escapes workspace: {path}")
    return resolved


# ---------------------------------------------------------------------------
# Tool implementations (synchronous)
# ---------------------------------------------------------------------------

def tool_web_search(query, config):
    """Web search via existing search_adapter."""
    try:
        from search_adapter import _search_via_tavily, _search_via_serpapi
    except ImportError:
        return json.dumps({"error": "search_adapter not available", "results": []})

    # 自动检测搜索提供商：优先用户配置，否则根据可用 API Key 判断
    provider = config.get("SEARCH_PROVIDER", "").lower()
    if not provider:
        if config.get("TAVILY_API_KEY"):
            provider = "tavily"
        elif config.get("SERPAPI_API_KEY"):
            provider = "serpapi"
        else:
            return json.dumps({"error": "No search API key configured", "results": []})
    try:
        if provider == "tavily":
            results = _search_via_tavily([query], config, fetch_top_n=5)
        elif provider == "serpapi":
            results = _search_via_serpapi([query], config, fetch_top_n=5)
        else:
            return json.dumps({"error": f"Unknown search provider: {provider}", "results": []})

        formatted = []
        for item in results:
            formatted.append({
                "title": item.get("title", ""),
                "url": item.get("url", ""),
                "snippet": item.get("content", "")[:500],
            })
        return json.dumps({"results": formatted}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e), "results": []}, ensure_ascii=False)


def tool_run_python(code, workspace):
    """Execute Python code. Subprocess in dev, exec() when frozen (PyInstaller)."""
    is_frozen = getattr(sys, 'frozen', False)

    if is_frozen:
        return _run_python_exec(code, workspace)
    else:
        return _run_python_subprocess(code, workspace)


def _run_python_subprocess(code, workspace):
    """Dev mode: subprocess.run python3."""
    try:
        result = subprocess.run(
            [sys.executable, "-c", code],
            cwd=workspace,
            capture_output=True,
            text=True,
            timeout=30,
            encoding="utf-8",
            env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"},
        )
        parts = []
        if result.stdout:
            parts.append(result.stdout)
        if result.stderr:
            parts.append(f"[stderr]\n{result.stderr}")
        output = "\n".join(parts).strip() or "(no output)"
        if len(output) > 10000:
            output = output[:10000] + "\n... [truncated]"
        return output
    except subprocess.TimeoutExpired:
        return "Error: Code execution timed out (30s)."
    except Exception as e:
        return f"Error: {e}"


def _run_python_exec(code, workspace):
    """PyInstaller mode: exec() with threading timeout."""
    result = {"output": "", "error": None}
    old_cwd = os.getcwd()

    def _exec():
        try:
            os.chdir(workspace)
            import io
            buf = io.StringIO()
            # 使用线程局部的 stdout 重定向，避免干扰主线程的 JSON 输出
            saved_stdout = sys.stdout
            saved_stderr = sys.stderr
            err_buf = io.StringIO()
            sys.stdout = buf
            sys.stderr = err_buf
            try:
                exec(code, {"__builtins__": __builtins__})
            finally:
                sys.stdout = saved_stdout
                sys.stderr = saved_stderr
            output = buf.getvalue()
            err_output = err_buf.getvalue()
            if err_output:
                output += f"\n[stderr]\n{err_output}"
            result["output"] = output
        except Exception as e:
            result["error"] = str(e)
        finally:
            os.chdir(old_cwd)

    t = threading.Thread(target=_exec)
    t.start()
    t.join(timeout=30)
    if t.is_alive():
        return "Error: Code execution timed out (30s)."
    if result["error"]:
        return f"Error: {result['error']}"
    output = result["output"].strip() or "(no output)"
    if len(output) > 10000:
        output = output[:10000] + "\n... [truncated]"
    return output


def tool_read_file(path, workspace):
    """Read file from workspace with path validation."""
    try:
        resolved = validate_path(path, workspace)
        if not os.path.exists(resolved):
            return f"Error: File not found: {path}"
        if not os.path.isfile(resolved):
            return f"Error: Not a file: {path}"
        size = os.path.getsize(resolved)
        if size > 20_000_000:
            return f"Error: File too large ({size} bytes)"

        ext = os.path.splitext(resolved)[1].lower()
        if ext in (".xlsx", ".xls"):
            return _read_excel(resolved)
        elif ext == ".docx":
            return _read_docx(resolved)
        elif ext == ".pdf":
            return _read_pdf(resolved)
        else:
            return _read_text(resolved)
    except ValueError as e:
        return f"Error: {e}"
    except Exception as e:
        return f"Error reading file: {e}"


def _read_text(path):
    """Read text file with encoding detection."""
    raw = open(path, "rb").read()
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            text = raw.decode(enc)
            if len(text) > 50000:
                text = text[:50000] + "\n... [truncated]"
            return text
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def _read_excel(path):
    """Extract Excel content as text."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for name in wb.sheetnames:
            ws = wb[name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                rows.append(" | ".join(cells))
                if len(rows) > 500:
                    rows.append("... (truncated)")
                    break
            parts.append(f"## Sheet: {name}\n" + "\n".join(rows))
        wb.close()
        result = "\n\n".join(parts)
        if len(result) > 50000:
            result = result[:50000] + "\n... [truncated]"
        return result
    except Exception as e:
        return f"Error reading Excel: {e}"


def _read_docx(path):
    """Extract Word document text."""
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        result = "\n".join(parts)
        if len(result) > 50000:
            result = result[:50000] + "\n... [truncated]"
        return result
    except Exception as e:
        return f"Error reading Word: {e}"


def _read_pdf(path):
    """Extract PDF text."""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages[:50]):
                text = page.extract_text()
                if text:
                    parts.append(f"--- Page {i+1} ---\n{text}")
        result = "\n\n".join(parts)
        if len(result) > 50000:
            result = result[:50000] + "\n... [truncated]"
        return result or "(empty PDF)"
    except Exception as e:
        return f"Error reading PDF: {e}"


def tool_write_file(path, content, workspace):
    """Write file to workspace with path validation."""
    try:
        resolved = validate_path(path, workspace)
        os.makedirs(os.path.dirname(resolved), exist_ok=True)
        with open(resolved, "w", encoding="utf-8") as f:
            f.write(content)
        return f"OK: wrote {len(content)} chars to {path}"
    except ValueError as e:
        return f"Error: {e}"
    except Exception as e:
        return f"Error writing file: {e}"


def execute_tool(name, args, workspace, config):
    """Tool dispatcher."""
    if name == "web_search":
        return tool_web_search(args.get("query", ""), config)
    elif name == "run_python":
        return tool_run_python(args.get("code", ""), workspace)
    elif name == "read_file":
        return tool_read_file(args.get("path", ""), workspace)
    elif name == "write_file":
        return tool_write_file(args.get("path", ""), args.get("content", ""), workspace)
    else:
        return f"Error: Unknown tool '{name}'"


# ---------------------------------------------------------------------------
# LLM calling with tools
# ---------------------------------------------------------------------------

def _resolve_provider(config):
    """Extract endpoint, api_key, model from config."""
    provider = config.get("LLM_PROVIDER", "deepseek").lower()
    if provider == "claude":
        # Agent 模式需要 function-calling，Claude CLI 不支持，回退到 deepseek
        logger.warning("Agent mode: Claude not supported, falling back to deepseek")
        provider = "deepseek"

    endpoint = PROVIDER_ENDPOINTS.get(provider)
    if not endpoint:
        raise ValueError(f"Unsupported provider for agent: {provider}")

    key_name = PROVIDER_KEY_NAMES[provider]
    api_key = config.get(key_name, "")
    if not api_key:
        raise ValueError(f"Missing API key: {key_name}")

    model_key, default_model = PROVIDER_MODEL_NAMES[provider]
    model = config.get(model_key, default_model)

    return endpoint, api_key, model


def call_llm_with_tools(messages, config, tools=None):
    """Call OpenAI-compatible API with function-calling support."""
    import requests

    endpoint, api_key, model = _resolve_provider(config)

    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 8192,
    }
    if tools:
        payload["tools"] = tools

    resp = requests.post(
        endpoint,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json=payload,
        timeout=300,
    )

    if resp.status_code != 200:
        raise RuntimeError(f"LLM API error: HTTP {resp.status_code} {resp.text[:300]}")

    return resp.json()


# ---------------------------------------------------------------------------
# Agent main loop
# ---------------------------------------------------------------------------

def run_agent_loop(topic, config, emit_fn, workspace,
                   template_prompt="", file_contents="",
                   file_formats=None, max_turns=15):
    """
    Run the multi-turn agent loop.

    Args:
        topic: User topic/task description
        config: Provider config dict
        emit_fn: Callback for progress events: emit_fn(type, **kwargs)
        workspace: Absolute path to workspace directory
        template_prompt: Template-specific prompt to merge into system prompt
        file_contents: Extracted text from uploaded files
        file_formats: List of dicts with file format info [{name, ext, path}]
        max_turns: Maximum agent turns

    Returns:
        HTML content string, or None if agent didn't produce output
    """
    from agent_prompts import get_agent_system_prompt

    # 替换模板中的 {{TOPIC}} 占位符
    if template_prompt and "{{TOPIC}}" in template_prompt:
        template_prompt = template_prompt.replace("{{TOPIC}}", topic)

    system_prompt = get_agent_system_prompt(template_prompt, file_formats)

    # Build user message
    user_content = f"请对「{topic}」进行深度调研和创作。\n\n"

    if file_contents:
        user_content += "## 用户上传的参考资料\n\n"
        user_content += file_contents + "\n\n"

    if file_formats:
        user_content += "## 上传文件信息\n\n"
        user_content += "以下文件已复制到 input/ 目录：\n"
        for f in file_formats:
            user_content += f"- `input/{f['name']}` (格式: .{f.get('ext', '?')})\n"
        user_content += "\n"

    user_content += (
        "请开始工作。使用工具进行调研和创作，"
        "最终将文章写入 output/article.html。"
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content},
    ]

    logger.info("Agent loop start: topic=%s, max_turns=%d", topic[:60], max_turns)

    for turn in range(max_turns):
        emit_fn("progress", stage="agent",
                message=f"Agent 第 {turn+1}/{max_turns} 轮")

        try:
            t0 = time.monotonic()
            response = call_llm_with_tools(messages, config, tools=TOOL_DEFINITIONS)
            elapsed = round((time.monotonic() - t0) * 1000)
        except Exception as e:
            logger.error("LLM call failed at turn %d: %s", turn+1, e)
            emit_fn("progress", stage="agent",
                    message=f"LLM 调用失败: {e}")
            break

        choices = response.get("choices", [])
        if not choices:
            emit_fn("progress", stage="agent", message="LLM 返回空响应")
            break

        message = choices[0].get("message", {})

        # Build assistant message for conversation history
        assistant_msg = {"role": "assistant"}
        assistant_msg["content"] = message.get("content") or ""
        if message.get("tool_calls"):
            assistant_msg["tool_calls"] = message["tool_calls"]
        messages.append(assistant_msg)

        # Log token usage
        usage = response.get("usage", {})
        if usage:
            logger.info("Turn %d: %dms, tokens in=%d out=%d",
                        turn+1, elapsed,
                        usage.get("prompt_tokens", 0),
                        usage.get("completion_tokens", 0))

        # Check for tool calls
        tool_calls = message.get("tool_calls")
        if not tool_calls:
            emit_fn("progress", stage="agent", message="Agent 完成创作")
            break

        # Execute each tool call
        for tc in tool_calls:
            tc_id = tc.get("id", "")
            func = tc.get("function", {})
            tool_name = func.get("name", "unknown")

            try:
                args = json.loads(func.get("arguments", "{}"))
            except json.JSONDecodeError:
                args = {}

            # Progress: tool call start
            preview = _tool_preview(tool_name, args)
            emit_fn("progress", stage="agent",
                    message=f"🔧 {tool_name}: {preview}")

            # Execute
            t0 = time.monotonic()
            result = execute_tool(tool_name, args, workspace, config)
            tool_ms = round((time.monotonic() - t0) * 1000)

            emit_fn("progress", stage="agent",
                    message=f"✓ {tool_name} 完成 ({tool_ms}ms)")

            # Append tool result
            messages.append({
                "role": "tool",
                "tool_call_id": tc_id,
                "content": result,
            })
    else:
        emit_fn("progress", stage="agent",
                message=f"Agent 达到最大轮次 ({max_turns})")

    # Read output files
    html_path = os.path.join(workspace, "output", "article.html")
    if os.path.exists(html_path):
        with open(html_path, "r", encoding="utf-8") as f:
            return f.read()

    # Fallback: check if LLM's last message contains HTML
    if messages and messages[-1].get("role") == "assistant":
        content = messages[-1].get("content", "")
        if "<section" in content or "<div" in content:
            return content

    return None


def _tool_preview(tool_name, args):
    """Short preview of tool call for progress display."""
    if tool_name == "web_search":
        return args.get("query", "")[:60]
    elif tool_name == "run_python":
        code = args.get("code", "")
        first_line = code.split("\n")[0][:60]
        return first_line
    elif tool_name == "read_file":
        return args.get("path", "")
    elif tool_name == "write_file":
        path = args.get("path", "")
        size = len(args.get("content", ""))
        return f"{path} ({size} chars)"
    return ""
